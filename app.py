import streamlit as st
import pandas as pd
from io import BytesIO
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import unicodedata
import re


# =========================================================
# CONFIGURACIÓN
# =========================================================

st.set_page_config(
    page_title="Generador EOD OXXO",
    page_icon="🟢",
    layout="wide"
)

st.title("🟢 Generador de Reporte EOD OXXO")

st.write(
    "Carga el CSV de encuestas, selecciona la tienda y diligencia "
    "la información que corresponde al análisis."
)


# =========================================================
# PLANTILLA
# =========================================================

# La plantilla debe estar en el mismo repositorio que app.py.
# Nombre exacto:
# Plantilla_EOD_OXXO.docx

RUTA_PLANTILLA = (
    Path(__file__).parent / "Plantilla_EOD_OXXO.docx"
)


# =========================================================
# FUNCIONES DE DATOS
# =========================================================

def normalizar_texto(texto):
    """
    Normaliza texto para encontrar columnas aunque tengan:
    - Mayúsculas/minúsculas diferentes
    - Tildes
    - Espacios
    - Guiones bajos
    - Signos de puntuación
    """

    texto = str(texto).strip().lower()

    texto = unicodedata.normalize(
        "NFKD",
        texto
    )

    texto = "".join(
        c for c in texto
        if not unicodedata.combining(c)
    )

    texto = re.sub(
        r"[^a-z0-9]+",
        "",
        texto
    )

    return texto


def buscar_columna(df, nombres):
    """
    Busca una columna aunque existan diferencias
    de mayúsculas, tildes, espacios, signos, etc.
    """

    columnas = {
        normalizar_texto(c): c
        for c in df.columns
    }

    for nombre in nombres:

        clave = normalizar_texto(nombre)

        if clave in columnas:
            return columnas[clave]

    return None


def porcentaje(cantidad, total):

    if total == 0:
        return 0

    return round(
        (cantidad / total) * 100
    )


def tabla_frecuencia(df, columna):

    if (
        columna is None
        or columna not in df.columns
    ):
        return pd.DataFrame(
            columns=[
                "Respuesta",
                "Cantidad",
                "%"
            ]
        )

    datos = (
        df[columna]
        .dropna()
        .astype(str)
        .str.strip()
    )

    datos = datos[
        datos != ""
    ]

    if len(datos) == 0:

        return pd.DataFrame(
            columns=[
                "Respuesta",
                "Cantidad",
                "%"
            ]
        )

    tabla = (
        datos
        .value_counts()
        .reset_index()
    )

    tabla.columns = [
        "Respuesta",
        "Cantidad"
    ]

    tabla["%"] = tabla["Cantidad"].apply(
        lambda x: porcentaje(
            x,
            len(datos)
        )
    )

    return tabla


def formato_fecha(fecha):

    if fecha is None or pd.isna(fecha):
        return ""

    return pd.to_datetime(
        fecha
    ).strftime("%d/%m/%Y")


def formato_fecha_hora(fecha):

    if fecha is None or pd.isna(fecha):
        return ""

    return pd.to_datetime(
        fecha
    ).strftime("%d/%m/%Y %H:%M")


# =========================================================
# FUNCIONES PARA WORD
# =========================================================

def obtener_todos_los_parrafos(doc):
    """
    Obtiene párrafos normales y párrafos dentro
    de tablas.
    """

    parrafos = []

    # Párrafos normales
    parrafos.extend(
        doc.paragraphs
    )

    # Párrafos dentro de tablas
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                parrafos.extend(
                    cell.paragraphs
                )

    return parrafos


def reemplazar_placeholder_en_parrafo(
    paragraph,
    placeholder,
    valor
):
    """
    Reemplaza un placeholder en un párrafo.
    """

    texto = paragraph.text

    if placeholder not in texto:
        return False

    # Primero intentamos reemplazar
    # dentro de los runs existentes.
    for run in paragraph.runs:

        if placeholder in run.text:

            run.text = run.text.replace(
                placeholder,
                str(valor)
            )

            return True

    # Si el placeholder está dividido entre
    # varios runs, reconstruimos el párrafo.
    texto_nuevo = texto.replace(
        placeholder,
        str(valor)
    )

    paragraph.clear()

    run = paragraph.add_run(
        texto_nuevo
    )

    return True


def reemplazar_en_parrafos(
    doc,
    reemplazos
):

    for paragraph in doc.paragraphs:

        for viejo, nuevo in reemplazos.items():

            reemplazar_placeholder_en_parrafo(
                paragraph,
                viejo,
                nuevo
            )


def reemplazar_en_tablas(
    doc,
    reemplazos
):

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    for viejo, nuevo in reemplazos.items():

                        reemplazar_placeholder_en_parrafo(
                            paragraph,
                            viejo,
                            nuevo
                        )


def reemplazar_placeholders(
    doc,
    reemplazos
):

    reemplazar_en_parrafos(
        doc,
        reemplazos
    )

    reemplazar_en_tablas(
        doc,
        reemplazos
    )


# =========================================================
# ESCRITURA EN CELDAS
# =========================================================

def escribir_celda(
    cell,
    texto,
    bold=False
):

    cell.text = ""

    paragraph = cell.paragraphs[0]

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(
        str(texto)
    )

    run.bold = bold

    run.font.size = Pt(8.5)


def limpiar_celda(cell):

    cell.text = ""


# =========================================================
# LLENAR TABLAS
# =========================================================

def llenar_tabla_frecuencia(
    table,
    tabla_datos,
    max_filas=None
):
    """
    Llena las filas que YA existen en la plantilla.
    No crea tablas nuevas.
    No cambia el diseño.
    """

    if max_filas is None:

        max_filas = max(
            len(table.rows) - 1,
            0
        )

    datos = tabla_datos.head(
        max_filas
    )

    for i in range(max_filas):

        fila = table.rows[
            i + 1
        ]

        if i < len(datos):

            dato = datos.iloc[i]

            escribir_celda(
                fila.cells[0],
                dato["Respuesta"]
            )

            escribir_celda(
                fila.cells[1],
                int(
                    dato["Cantidad"]
                )
            )

            escribir_celda(
                fila.cells[2],
                f"{int(dato['%'])}%"
            )

        else:

            for cell in fila.cells:

                limpiar_celda(
                    cell
                )


def identificar_tabla_por_encabezado(
    table
):
    """
    Devuelve los encabezados de una tabla
    normalizados.
    """

    if not table.rows:
        return []

    return [
        normalizar_texto(
            cell.text
        )
        for cell in table.rows[0].cells
    ]


def llenar_datos_de_tablas(
    doc,
    tablas
):
    """
    Identifica las tablas por sus encabezados.
    """

    contador_lugar = 0

    for table in doc.tables:

        encabezados = identificar_tabla_por_encabezado(
            table
        )

        texto_tabla = " ".join(
            encabezados
        )

        # -------------------------------------------------
        # ALTERNATIVA DE COMPRA
        # -------------------------------------------------

        if (
            "competidor" in texto_tabla
            and "cantidad" in texto_tabla
            and "%" in texto_tabla
        ):

            llenar_tabla_frecuencia(
                table,
                tablas["alternativa"]
            )

        # -------------------------------------------------
        # OCUPACIÓN
        # -------------------------------------------------

        elif (
            "ocupacion" in texto_tabla
            and "cantidad" in texto_tabla
            and "%" in texto_tabla
        ):

            llenar_tabla_frecuencia(
                table,
                tablas["ocupacion"]
            )

        # -------------------------------------------------
        # MOTIVO
        # -------------------------------------------------

        elif (
            "motivo" in texto_tabla
            and "cantidad" in texto_tabla
            and "%" in texto_tabla
        ):

            llenar_tabla_frecuencia(
                table,
                tablas["motivo"]
            )

        # -------------------------------------------------
        # TRANSPORTE
        # -------------------------------------------------

        elif (
            (
                "mediotransporte" in texto_tabla
                or "transporte" in texto_tabla
            )
            and "cantidad" in texto_tabla
            and "%" in texto_tabla
        ):

            llenar_tabla_frecuencia(
                table,
                tablas["transporte"]
            )

        # -------------------------------------------------
        # ORIGEN / DESTINO
        # -------------------------------------------------

        elif (
            "lugar" in texto_tabla
            and "cantidad" in texto_tabla
            and "%" in texto_tabla
        ):

            if contador_lugar == 0:

                llenar_tabla_frecuencia(
                    table,
                    tablas["origen"]
                )

            else:

                llenar_tabla_frecuencia(
                    table,
                    tablas["destino"]
                )

            contador_lugar += 1


# =========================================================
# IMÁGENES
# =========================================================

def copiar_imagen(imagen):

    """
    Convierte la imagen subida por Streamlit
    en un BytesIO independiente.

    Esto evita problemas al reutilizar
    UploadedFile dentro de Word.
    """

    if imagen is None:
        return None

    try:

        contenido = imagen.getvalue()

        buffer = BytesIO(
            contenido
        )

        buffer.seek(0)

        return buffer

    except Exception:

        try:

            imagen.seek(0)

            contenido = imagen.read()

            buffer = BytesIO(
                contenido
            )

            buffer.seek(0)

            return buffer

        except Exception:

            return None


def poner_imagen_en_celda(
    cell,
    imagen,
    ancho=5.8
):
    """
    Inserta una imagen dentro de una celda.
    """

    # Limpiamos únicamente el contenido
    # de la celda.
    cell.text = ""

    paragraph = cell.paragraphs[0]

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    if imagen is None:

        return

    try:

        imagen.seek(0)

        run = paragraph.add_run()

        run.add_picture(
            imagen,
            width=Inches(ancho)
        )

    except Exception as e:

        paragraph.text = (
            "No se pudo insertar la imagen."
        )


def poner_imagen_por_placeholder(
    doc,
    placeholder,
    imagen,
    ancho=5.8
):
    """
    Busca un placeholder dentro de una celda
    y coloca la imagen allí.
    """

    if imagen is None:
        return False

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                if placeholder in cell.text:

                    imagen_buffer = copiar_imagen(
                        imagen
                    )

                    poner_imagen_en_celda(
                        cell,
                        imagen_buffer,
                        ancho
                    )

                    return True

    # También permite placeholder
    # en párrafos normales.
    for paragraph in doc.paragraphs:

        if placeholder in paragraph.text:

            try:

                paragraph.clear()

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )

                imagen_buffer = copiar_imagen(
                    imagen
                )

                if imagen_buffer is not None:

                    run = paragraph.add_run()

                    run.add_picture(
                        imagen_buffer,
                        width=Inches(ancho)
                    )

                return True

            except Exception:

                return False

    return False


# =========================================================
# FOTO DE FACHADA
# =========================================================

def insertar_foto_fachada(
    doc,
    imagen
):
    """
    Busca [FOTO_FACHADA] en la plantilla
    y lo reemplaza por la foto.

    IMPORTANTE:
    El placeholder debe estar ubicado debajo
    del nombre de la tienda en la plantilla.
    """

    if imagen is None:
        return False

    return poner_imagen_por_placeholder(
        doc,
        "[FOTO_FACHADA]",
        imagen,
        ancho=3.5
    )


# =========================================================
# REGISTRO FOTOGRÁFICO
# =========================================================

def llenar_registro_fotografico(
    doc,
    fotos
):
    """
    Busca los placeholders:

    [FOTO_1]
    [FOTO_2]
    [FOTO_3]
    [FOTO_4]

    y los reemplaza con las fotografías.
    """

    placeholders = [
        "[FOTO_1]",
        "[FOTO_2]",
        "[FOTO_3]",
        "[FOTO_4]"
    ]

    for i, placeholder in enumerate(
        placeholders
    ):

        if i < len(fotos):

            poner_imagen_por_placeholder(
                doc,
                placeholder,
                fotos[i],
                ancho=2.65
            )

        else:

            # Si no hay foto, eliminamos
            # el placeholder.
            reemplazar_placeholders(
                doc,
                {
                    placeholder: ""
                }
            )


# =========================================================
# INSIGHTS
# =========================================================

def llenar_insights(
    doc,
    insights
):

    placeholders = [
        "[INSIGHT_1]",
        "[INSIGHT_2]",
        "[INSIGHT_3]"
    ]

    reemplazos = {}

    for placeholder, texto in zip(
        placeholders,
        insights
    ):

        reemplazos[
            placeholder
        ] = texto

    reemplazar_placeholders(
        doc,
        reemplazos
    )


# =========================================================
# CREAR WORD
# =========================================================

def crear_word(
    datos,
    manuales,
    tablas
):

    if not RUTA_PLANTILLA.exists():

        raise FileNotFoundError(
            "No se encontró la plantilla "
            "'Plantilla_EOD_OXXO.docx'. "
            "Verifica que esté en el mismo "
            "repositorio que app.py."
        )

    # Abrimos la plantilla ORIGINAL
    doc = Document(
        str(RUTA_PLANTILLA)
    )

    # -----------------------------------------------------
    # DATOS PRINCIPALES
    # -----------------------------------------------------

    reemplazos = {

        "[NOMBRE TIENDA]":
            datos["tienda"],

        "[PERIODO]":
            datos["periodo"],

        "[FECHA_INICIO]":
            datos["inicio"],

        "[FECHA_FIN]":
            datos["finalizacion"],

        "[N_ENCUESTAS]":
            datos["encuestas"],

        "[DIA_MAS]":
            datos["dia_mas_encuestas"],

        "[CANTIDAD_DIA_MAS]":
            datos["cantidad_dia_mas"],

        "[EDAD_PROMEDIO]":
            datos["edad_promedio"],

        "[HOMBRES]":
            datos["hombres"],

        "[MUJERES]":
            datos["mujeres"],

        "[ESTRATO_PRINCIPAL]":
            datos["estrato"],

        "[ESTRATO_PORCENTAJE]":
            datos["estrato_porcentaje"],

        "[OCUPACION_PRINCIPAL]":
            datos["ocupacion"],

        "[PERCEPCION_SERVICIO]":
            manuales["percepcion"],

        "[RADIO_INFLUENCIA]":
            manuales["radio"]
    }

    reemplazar_placeholders(
        doc,
        reemplazos
    )

    # -----------------------------------------------------
    # TABLAS
    # -----------------------------------------------------

    llenar_datos_de_tablas(
        doc,
        tablas
    )

    # -----------------------------------------------------
    # INSIGHTS
    # -----------------------------------------------------

    llenar_insights(
        doc,
        [
            manuales["insight1"],
            manuales["insight2"],
            manuales["insight3"]
        ]
    )

    # -----------------------------------------------------
    # FOTO FACHADA
    # -----------------------------------------------------

    insertar_foto_fachada(
        doc,
        manuales["foto_fachada"]
    )

    # -----------------------------------------------------
    # FOTO RADIO
    # -----------------------------------------------------

    poner_imagen_por_placeholder(
        doc,
        "[PEGAR AQUÍ / IMAGEN DEL RADIO DE INFLUENCIA]",
        manuales["imagen_radio"],
        ancho=5.8
    )

    # -----------------------------------------------------
    # FOTO ISÓCRONA
    # -----------------------------------------------------

    poner_imagen_por_placeholder(
        doc,
        "[PEGAR AQUÍ / IMAGEN DE LA ISÓCRONA]",
        manuales["imagen_isocrona"],
        ancho=5.8
    )

    # -----------------------------------------------------
    # REGISTRO FOTOGRÁFICO
    # -----------------------------------------------------

    llenar_registro_fotografico(
        doc,
        manuales["fotos"]
    )

    # -----------------------------------------------------
    # GUARDAR
    # -----------------------------------------------------

    buffer = BytesIO()

    doc.save(
        buffer
    )

    buffer.seek(0)

    return buffer


# =========================================================
# 1. CARGAR CSV
# =========================================================

st.header(
    "1. Cargar archivo de encuestas"
)

archivo = st.file_uploader(
    "Selecciona el CSV de encuestas",
    type=["csv"]
)

if archivo is None:

    st.info(
        "Primero carga el archivo EOD_OXXO."
    )

    st.stop()


# =========================================================
# LEER CSV
# =========================================================

try:

    df = pd.read_csv(
        archivo,
        encoding="utf-8"
    )

except Exception:

    archivo.seek(0)

    df = pd.read_csv(
        archivo,
        encoding="latin-1"
    )


df.columns = [
    str(c).strip()
    for c in df.columns
]


st.success(
    f"Archivo cargado correctamente: "
    f"{len(df)} encuestas."
)


# =========================================================
# IDENTIFICAR COLUMNAS
# =========================================================

col_fecha = buscar_columna(
    df,
    [
        "CreationDate",
        "Creation Date"
    ]
)


col_tienda = buscar_columna(
    df,
    [
        "Nombre tienda estudiada"
    ]
)


col_edad = buscar_columna(
    df,
    [
        "Edad"
    ]
)


col_genero = buscar_columna(
    df,
    [
        "Género",
        "Genero"
    ]
)


col_estrato = buscar_columna(
    df,
    [
        "Estrato"
    ]
)


col_ocupacion = buscar_columna(
    df,
    [
        "Ocupación actual",
        "Ocupacion actual"
    ]
)


col_motivo = buscar_columna(
    df,
    [
        "Por qué compraría ahí?",
        "Por que compraria ahi?"
    ]
)


col_transporte = buscar_columna(
    df,
    [
        "Medio de transporte usado para llegar a OXXO",
        "Medio de transporte usado para llegar a OXXO?"
    ]
)


col_origen = buscar_columna(
    df,
    [
        "De dónde viene?",
        "De donde viene?"
    ]
)


col_destino = buscar_columna(
    df,
    [
        "Hacia dónde se dirige?",
        "Hacia donde se dirige?"
    ]
)


# =========================================================
# ALTERNATIVA DE COMPRA
# =========================================================

col_alternativa = buscar_columna(
    df,
    [

        "Dónde compraría sino es en OXXO?",

        "Dónde compraría si no es en OXXO?",

        "Dónde compraría si no fuera en OXXO?",

        "Donde compraria sino es en OXXO?",

        "Donde compraria si no es en OXXO?",

        "Donde compraria si no fuera en OXXO?",

        "Dónde compraría si no fuera OXXO?",

        "Donde compraria si no fuera OXXO?"
    ]
)


# =========================================================
# VALIDACIÓN TIENDA
# =========================================================

if col_tienda is None:

    st.error(
        "No encontré la columna "
        "'Nombre tienda estudiada'."
    )

    st.write(
        "Columnas encontradas:"
    )

    st.write(
        list(df.columns)
    )

    st.stop()


# =========================================================
# 2. SELECCIONAR TIENDA
# =========================================================

st.header(
    "2. Seleccionar tienda"
)

tiendas = (
    df[col_tienda]
    .dropna()
    .astype(str)
    .str.strip()
    .unique()
)

tiendas = sorted(
    tiendas
)


tienda_seleccionada = st.selectbox(
    "Selecciona el CR / tienda",
    tiendas
)


# =========================================================
# FOTO DE FACHADA
# =========================================================

st.subheader(
    "Foto de fachada"
)

st.write(
    "Sube la fotografía de la fachada correspondiente "
    "al CR seleccionado."
)

foto_fachada = st.file_uploader(
    "Subir foto de fachada",
    type=[
        "jpg",
        "jpeg",
        "png"
    ],
    key="foto_fachada"
)


# =========================================================
# FILTRAR TIENDA
# =========================================================

datos_tienda = df[
    df[col_tienda]
    .astype(str)
    .str.strip()
    ==
    tienda_seleccionada
].copy()


# =========================================================
# FECHAS
# =========================================================

fecha_inicio = None
fecha_final = None

dia_mas_encuestas = ""

cantidad_dia_mas = 0

dias_sin_encuestas = []


if col_fecha is not None:

    datos_tienda[col_fecha] = pd.to_datetime(
        datos_tienda[col_fecha],
        errors="coerce"
    )

    fechas = (
        datos_tienda[col_fecha]
        .dropna()
    )

    if len(fechas) > 0:

        fecha_inicio = fechas.min()

        fecha_final = fechas.max()

        conteo_dias = (
            fechas.dt.date
            .value_counts()
            .sort_index()
        )

        if len(conteo_dias) > 0:

            dia_mas_encuestas = (
                conteo_dias.idxmax()
            )

            cantidad_dia_mas = int(
                conteo_dias.max()
            )

            todos_los_dias = pd.date_range(
                start=fecha_inicio.date(),
                end=fecha_final.date(),
                freq="D"
            ).date

            dias_sin_encuestas = [

                dia.strftime(
                    "%d/%m/%Y"
                )

                for dia in todos_los_dias

                if dia not in conteo_dias.index
            ]


# =========================================================
# 3. INFORMACIÓN AUTOMÁTICA
# =========================================================

st.header(
    "3. Información automática"
)

total_encuestas = len(
    datos_tienda
)


if (
    col_fecha is not None
    and fecha_inicio is not None
):

    if (
        fecha_inicio.date()
        ==
        fecha_final.date()
    ):

        periodo = formato_fecha(
            fecha_inicio
        )

    else:

        periodo = (
            f"{formato_fecha(fecha_inicio)} - "
            f"{formato_fecha(fecha_final)}"
        )

else:

    periodo = ""


col1, col2, col3 = st.columns(3)


with col1:

    st.metric(
        "Encuestas",
        total_encuestas
    )


with col2:

    st.metric(
        "Inicio",
        formato_fecha_hora(
            fecha_inicio
        )
    )


with col3:

    st.metric(
        "Finalización",
        formato_fecha_hora(
            fecha_final
        )
    )


col4, col5, col6 = st.columns(3)


with col4:

    st.metric(
        "Día con más encuestas",
        str(
            dia_mas_encuestas
        )
    )


with col5:

    st.metric(
        "Cantidad ese día",
        cantidad_dia_mas
    )


with col6:

    st.write(
        "**Periodo**"
    )

    st.write(
        periodo
    )


if dias_sin_encuestas:

    st.warning(
        "Días sin encuestas: "
        +
        ", ".join(
            dias_sin_encuestas
        )
    )

else:

    st.success(
        "No se encontraron días sin encuestas."
    )


# =========================================================
# 4. PERFIL DEL CLIENTE
# =========================================================

st.header(
    "4. Perfil del cliente"
)


edad_promedio = ""


if col_edad is not None:

    edades = pd.to_numeric(
        datos_tienda[col_edad],
        errors="coerce"
    )

    if edades.notna().any():

        edad_promedio = round(
            edades.mean(),
            1
        )


hombres = 0
mujeres = 0


if col_genero is not None:

    genero = (
        datos_tienda[col_genero]
        .astype(str)
        .str.strip()
        .str.lower()
    )

    hombres = int(
        genero.str.contains(
            "hombre",
            na=False
        ).sum()
    )

    mujeres = int(
        genero.str.contains(
            "mujer",
            na=False
        ).sum()
    )


estrato_tabla = tabla_frecuencia(
    datos_tienda,
    col_estrato
)


ocupacion_tabla = tabla_frecuencia(
    datos_tienda,
    col_ocupacion
)


if not estrato_tabla.empty:

    estrato_principal = (
        estrato_tabla.iloc[0]["Respuesta"]
    )

    estrato_porcentaje = (
        estrato_tabla.iloc[0]["%"]
    )

else:

    estrato_principal = "N/D"

    estrato_porcentaje = 0


p1, p2, p3, p4 = st.columns(4)


with p1:

    st.metric(
        "Edad promedio",
        (
            f"{edad_promedio} años"
            if edad_promedio != ""
            else "N/D"
        )
    )


with p2:

    st.metric(
        "Hombres",
        hombres
    )


with p3:

    st.metric(
        "Mujeres",
        mujeres
    )


with p4:

    st.metric(
        "Estrato principal",
        (
            f"{estrato_principal} "
            f"({estrato_porcentaje}%)"
        )
    )


if not ocupacion_tabla.empty:

    st.write(
        "**Ocupación principal**"
    )

    st.dataframe(
        ocupacion_tabla,
        use_container_width=True,
        hide_index=True
    )


# =========================================================
# 5. INFORMACIÓN DEL ESTUDIO
# =========================================================

st.header(
    "5. Información del estudio"
)


tablas = {

    "ocupacion":
        ocupacion_tabla,

    "motivo":
        tabla_frecuencia(
            datos_tienda,
            col_motivo
        ),

    "transporte":
        tabla_frecuencia(
            datos_tienda,
            col_transporte
        ),

    "origen":
        tabla_frecuencia(
            datos_tienda,
            col_origen
        ),

    "destino":
        tabla_frecuencia(
            datos_tienda,
            col_destino
        ),

    "alternativa":
        tabla_frecuencia(
            datos_tienda,
            col_alternativa
        )
}


# =========================================================
# MOSTRAR SI ENCONTRÓ ALTERNATIVA
# =========================================================

if col_alternativa is None:

    st.warning(
        "No se identificó automáticamente la columna "
        "de alternativa de compra."
    )

    st.write(
        "Columnas disponibles:"
    )

    st.write(
        list(df.columns)
    )

else:

    st.success(
        f"Columna de alternativa identificada: "
        f"'{col_alternativa}'"
    )


# =========================================================
# TABLAS EN STREAMLIT
# =========================================================

t1, t2 = st.columns(2)


with t1:

    st.subheader(
        "Principal motivo de compra"
    )

    st.dataframe(
        tablas["motivo"],
        use_container_width=True,
        hide_index=True
    )


with t2:

    st.subheader(
        "Medio de llegada"
    )

    st.dataframe(
        tablas["transporte"],
        use_container_width=True,
        hide_index=True
    )


t3, t4 = st.columns(2)


with t3:

    st.subheader(
        "De dónde viene"
    )

    st.dataframe(
        tablas["origen"],
        use_container_width=True,
        hide_index=True
    )


with t4:

    st.subheader(
        "A dónde se dirige"
    )

    st.dataframe(
        tablas["destino"],
        use_container_width=True,
        hide_index=True
    )


st.subheader(
    "Alternativa de compra"
)

st.dataframe(
    tablas["alternativa"],
    use_container_width=True,
    hide_index=True
)


# =========================================================
# 6. INFORMACIÓN MANUAL
# =========================================================

st.header(
    "6. Información manual"
)

st.info(
    "Diligencia únicamente la información "
    "que corresponde al análisis de la tienda."
)


percepcion = st.text_area(
    "Percepción del servicio",
    placeholder=(
        "Escribe aquí la percepción "
        "del servicio..."
    )
)


radio = st.text_input(
    "Radio de influencia",
    placeholder=(
        "Ejemplo: 100m: 47 | "
        "200m: 48 | "
        "300m: 74 | "
        "+300m: 67"
    )
)


# =========================================================
# IMAGEN RADIO
# =========================================================

imagen_radio = st.file_uploader(
    "Pantallazo del radio de influencia",
    type=[
        "jpg",
        "jpeg",
        "png"
    ],
    key="imagen_radio"
)


# =========================================================
# IMAGEN ISÓCRONA
# =========================================================

imagen_isocrona = st.file_uploader(
    "Pantallazo de la isócrona",
    type=[
        "jpg",
        "jpeg",
        "png"
    ],
    key="imagen_isocrona"
)


# =========================================================
# INSIGHTS
# =========================================================

st.subheader(
    "Insights clave"
)


insight1 = st.text_area(
    "Insight 1",
    placeholder=(
        "Escribe el primer insight..."
    )
)


insight2 = st.text_area(
    "Insight 2",
    placeholder=(
        "Escribe el segundo insight..."
    )
)


insight3 = st.text_area(
    "Insight 3",
    placeholder=(
        "Escribe el tercer insight..."
    )
)


# =========================================================
# 7. REGISTRO FOTOGRÁFICO
# =========================================================

st.header(
    "7. Registro fotográfico"
)


fotos = st.file_uploader(
    "Sube las fotografías de la tienda "
    "(máximo 4)",
    type=[
        "jpg",
        "jpeg",
        "png"
    ],
    accept_multiple_files=True,
    key="fotos"
)


if len(fotos) > 4:

    st.warning(
        "La plantilla tiene 4 espacios para "
        "fotografías. Solo se utilizarán "
        "las primeras 4."
    )

    fotos = fotos[:4]


# =========================================================
# PREPARAR DATOS
# =========================================================

datos = {

    "tienda":
        tienda_seleccionada,

    "encuestas":
        total_encuestas,

    "periodo":
        periodo,

    "inicio":
        formato_fecha_hora(
            fecha_inicio
        ),

    "finalizacion":
        formato_fecha_hora(
            fecha_final
        ),

    "dia_mas_encuestas":
        str(
            dia_mas_encuestas
        ),

    "cantidad_dia_mas":
        cantidad_dia_mas,

    "edad_promedio":
        edad_promedio,

    "hombres":
        hombres,

    "mujeres":
        mujeres,

    "estrato":
        estrato_principal,

    "estrato_porcentaje":
        estrato_porcentaje,

    "ocupacion":
        (
            ocupacion_tabla.iloc[0]["Respuesta"]
            if not ocupacion_tabla.empty
            else "N/D"
        )
}


manuales = {

    "percepcion":
        percepcion,

    "radio":
        radio,

    "foto_fachada":
        foto_fachada,

    "imagen_radio":
        imagen_radio,

    "imagen_isocrona":
        imagen_isocrona,

    "insight1":
        insight1,

    "insight2":
        insight2,

    "insight3":
        insight3,

    "fotos":
        fotos
}


# =========================================================
# 8. GENERAR REPORTE
# =========================================================

st.header(
    "8. Generar reporte"
)


if st.button(
    "🟢 GENERAR REPORTE",
    type="primary"
):

    try:

        documento = crear_word(
            datos,
            manuales,
            tablas
        )

        nombre_archivo = (
            "Reporte_EOD_"
            f"{tienda_seleccionada}.docx"
        )

        st.success(
            "¡Reporte generado correctamente! 🎉"
        )

        st.download_button(
            label="📥 Descargar reporte Word",
            data=documento,
            file_name=nombre_archivo,
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            )
        )

    except Exception as e:

        st.error(
            "No se pudo generar el reporte."
        )

        st.exception(e)
